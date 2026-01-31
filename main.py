import os
import subprocess
import shutil
import csv
from pathlib import Path
from datetime import datetime
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
  
load_dotenv()  
  
GPT_MODEL = os.getenv("GPT_MODEL", "gpt-4.1-mini")  
AI_ENDPOINT = os.getenv("AI_ENDPOINT", "https://api.openai.com/v1")  
TEMPLATE_CV = Path(os.getenv("TEMPLATE_CV", "Template/CV.pdf"))  
TEMPLATE_COVER_LETTER = Path(os.getenv("TEMPLATE_COVER_LETTER_EXAMPLE", "Template/Cover Letter.docx"))  
OUTPUT_BASE_DIR = Path(os.getenv("OUTPUT_BASE_DIR", "."))  
COVER_LETTER_PROMPT_BASE = os.getenv("COVER_LETTER_PROMPT", "Generate a cover letter based on the example.")  
TRACKER_FILE = Path(os.getenv("TRACKER_FILE", "list.csv"))  
  
CSV_FIELDNAMES = ['url', 'status', 'company', 'position', 'location', 'language', 'date_processed', 'notes']  

def get_ai_client():
    return OpenAI(api_key=os.getenv("OPENAI_API_KEY"), base_url=AI_ENDPOINT)
  
def read_csv_entries() -> list[dict]:  
    if not TRACKER_FILE.exists():  
        return []  
    with open(TRACKER_FILE, 'r', newline='', encoding='utf-8-sig') as f:  
        reader = csv.DictReader(f)  
        return list(reader)  
  
def write_csv_entries(entries: list[dict]):  
    cleaned_entries = []  
    for entry in entries:  
        cleaned = {k: v for k, v in entry.items() if k and k in CSV_FIELDNAMES}  
        for field in CSV_FIELDNAMES:  
            if field not in cleaned:  
                cleaned[field] = ''  
        cleaned_entries.append(cleaned)  
      
    with open(TRACKER_FILE, 'w', newline='', encoding='utf-8-sig') as f:  
        writer = csv.DictWriter(f, fieldnames=CSV_FIELDNAMES)  
        writer.writeheader()  
        writer.writerows(cleaned_entries)  
  
def update_csv_entry(url: str, updates: dict):  
    entries = read_csv_entries()  
    for entry in entries:  
        if entry['url'] == url:  
            entry.update(updates)  
            break  
    write_csv_entries(entries)  
  
def fetch_job_posting_from_url(url: str) -> str:  
    import ollama  
    client = ollama.Client()  
    client._client.headers['Authorization'] = f'Bearer {os.getenv("OLLAMA_API_KEY")}'  
    result = client.web_fetch(url=url)  
    return result.content  
  
def get_company_name(job_posting: str) -> str:  
    client = get_ai_client()  
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": "Extract only the company name. Respond with the name only, nothing else."},  
            {"role": "user", "content": job_posting}  
        ],  
        temperature=0  
    )  
    return response.choices[0].message.content.strip()  
  
def get_position_title(job_posting: str) -> str:  
    client = get_ai_client()  
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": "Extract only the job position/title. Respond with the title only, nothing else."},  
            {"role": "user", "content": job_posting}  
        ],  
        temperature=0  
    )  
    return response.choices[0].message.content.strip()  
  
def get_job_location(job_posting: str) -> str:  
    client = get_ai_client()  
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": "Extract only the job location (city, country or remote). Respond with the location only, nothing else. If not specified, respond with 'Not specified'."},  
            {"role": "user", "content": job_posting}  
        ],  
        temperature=0  
    )  
    return response.choices[0].message.content.strip()  
  
def detect_language(job_posting: str) -> str:  
    client = get_ai_client()  
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": "Detect the language of the text and respond with only the ISO 639-1 language code (e.g., 'fr', 'en', 'de', 'es', 'ja', etc.). Return ONLY the 2-letter code, nothing else."},  
            {"role": "user", "content": job_posting[:500]}  
        ],  
        temperature=0  
    )  
    return response.choices[0].message.content.strip().lower()  
  
def get_cover_letter_label(language: str) -> str:  
    client = get_ai_client()  
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": f"Return only the translation of 'Cover Letter' in the language with ISO code '{language}'. Return ONLY the translation, nothing else."},  
            {"role": "user", "content": "Cover Letter"}  
        ],  
        temperature=0  
    )  
    return response.choices[0].message.content.strip()  
  
def get_email_subject(language: str) -> str:  
    client = get_ai_client()  
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": f"Return only the translation of 'Application' (job application) in the language with ISO code '{language}'. Return ONLY the translation, nothing else."},  
            {"role": "user", "content": "Application"}  
        ],  
        temperature=0  
    )  
    return response.choices[0].message.content.strip()  
  
def get_cv_updates(job_posting: str, language: str) -> dict:  
    client = get_ai_client()  
    todays_date = datetime.now().strftime("%Y-%m-%d")  
      
    prompt = f"""Extract CV header updates from this job posting.  
JOB POSTING: {job_posting}  
TODAY'S DATE: {todays_date}  
Return JSON with:  
- "internship_field": The SECTOR/DOMAIN of the internship, NOT the job title. Use 1-2 words maximum. Think about what broad business function or department this role belongs to. Examples:  
  * Financial Analyst → "Finance"  
  * Management Controller → "Controlling"  
  * Digital Marketing Assistant → "Marketing"  
  Identify the core business function, not the specific role. Keep it VERY SHORT (1-2 words) and in the SAME LANGUAGE as the job posting (ISO code: {language}).  
- "start_date": Start month in the format and language of the job posting (e.g., "Février", "March", "März"). If not specified, use the current month+1 in the appropriate language.  
Return ONLY valid JSON."""  
  
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": "Return only valid JSON. No other text."},  
            {"role": "user", "content": prompt}  
        ],  
        temperature=0  
    )  
      
    import json  
    content = response.choices[0].message.content.strip()  
    content = content.replace('```json', '').replace('```', '').strip()  
      
    try:  
        return json.loads(content)  
    except json.JSONDecodeError:  
        print(f"Warning: Invalid JSON from GPT: {content}")  
        return {"internship_field": "Finance", "start_date": "February"}  
  
def get_new_cover_letter_paragraphs(job_posting: str, company_name: str, template_path: Path, language: str) -> list:  
    client = get_ai_client()  
    doc = Document(template_path)  
    original_paragraphs = [p.text for p in doc.paragraphs]  
    todays_date = datetime.now().strftime("%Y-%m-%d")  
      
    prompt = f"""{COVER_LETTER_PROMPT_BASE}  
Return a JSON array of paragraph texts (one string per paragraph) matching the EXACT structure of the example.  
IMPORTANT: The cover letter MUST be written in the SAME LANGUAGE as the job posting (ISO code: {language}).  
JOB POSTING: {job_posting}  
COMPANY: {company_name}  
TODAY'S DATE: {todays_date}  
ORIGINAL PARAGRAPHS:  
{chr(10).join(f'{i}: {p}' for i, p in enumerate(original_paragraphs))}  
Return a JSON array with the same number of elements, where each element is the new text for that paragraph in the appropriate language."""  
      
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": "Return only a valid JSON array of strings. No other text."},  
            {"role": "user", "content": prompt}  
        ],  
        temperature=0.5  
    )  
      
    import json  
    content = response.choices[0].message.content.strip()  
    content = content.replace('```json', '').replace('```', '').strip()  
      
    try:  
        return json.loads(content)  
    except json.JSONDecodeError as e:  
        print(f"JSON decode error: {e}")  
        print(f"Response content:\n{content[:500]}")  
        import re  
        content = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', content)  
        try:  
            return json.loads(content)  
        except json.JSONDecodeError:  
            print("Failed to parse even after cleanup. Returning original paragraphs.")  
            return original_paragraphs  
  
def modify_cv_header(cv_path: Path, output_path: Path, updates: dict):  
    shutil.copy(cv_path, output_path)  
    doc = Document(output_path)  
      
    for table in doc.tables:  
        for row in table.rows:  
            for cell in row.cells:  
                for para in cell.paragraphs:  
                    for run in para.runs:  
                        if any(field in run.text for field in ["Finance", "Contrôle", "Data", "Marketing", "Audit", "Gestion", "Risk", "Analysis"]):  
                            for field in ["Finance", "Contrôle", "Data", "Marketing", "Audit", "Gestion", "Risk", "Analysis"]:  
                                run.text = run.text.replace(field, updates['internship_field'])  
                        for month in ["January", "February", "March", "April", "May", "June",  
                                     "Janvier", "Février", "Mars", "Avril", "Mai", "Juin",  
                                     "Januar", "Februar", "März", "April", "Mai", "Juni"]:  
                            run.text = run.text.replace(month, updates['start_date'])  
      
    doc.save(output_path)  
  
def modify_cover_letter(template_path: Path, output_path: Path, new_paragraphs: list):  
    shutil.copy(template_path, output_path)  
    doc = Document(output_path)  
      
    for i, para in enumerate(doc.paragraphs):  
        if i < len(new_paragraphs):  
            for run in para.runs[1:]:  
                run.text = ""  
            if para.runs:  
                para.runs[0].text = new_paragraphs[i]  
            else:  
                para.text = new_paragraphs[i]  
      
    doc.save(output_path)  
  
def convert_to_pdf_linux(docx_path: Path, pdf_path: Path):  
    try:  
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', str(pdf_path.parent), str(docx_path)], check=True, capture_output=True)  
    except subprocess.CalledProcessError as e:  
        print(f"Warning: PDF conversion error: {e}")  
  
def get_unique_filename(base_path: Path) -> Path:  
    if not base_path.exists():  
        return base_path  
    counter = 2  
    while True:  
        new_path = base_path.parent / f"{base_path.stem}-{counter}{base_path.suffix}"  
        if not new_path.exists():  
            return new_path  
        counter += 1  
  
def generate_email_body(company_name: str, position_title: str, language: str, template_path: Path) -> str:  
    client = get_ai_client()  
      
    doc = Document(template_path)  
    candidate_name = doc.paragraphs[0].text.strip()  
      
    phone_number = ""  
    for para in doc.paragraphs[:5]:  
        text = para.text  
        import re  
        phone_match = re.search(r'\+?\d[\d\s\-\(\)]{8,}', text)  
        if phone_match:  
            phone_number = phone_match.group(0)  
            break  
      
    prompt = f"""Generate a professional, natural email body for a job application.  
  
COMPANY: {company_name}  
POSITION: {position_title}  
LANGUAGE: ISO code {language}  
  
The email should:  
- Use natural, conversational language (avoid overly formal phrases like "Je reste à votre disposition", "Veuillez trouver ci-joint", etc.)  
- Be short but show genuine interest (3-4 sentences)  
- Mention the specific position being applied for  
- Briefly mention 1-2 relevant qualifications or interests that connect to the role  
- Mention that CV and cover letter are attached  
- Include proper line breaks between greeting, body paragraphs, and closing  
- Be written ENTIRELY in the language corresponding to ISO code {language}  
  
AVOID:  
- Overly stiff formal phrases  
- Generic templates  
- Placeholder text like [Your Name]  
  
The signature will be added separately, so end with just the closing phrase (e.g., "Cordialement," or "Best regards,")  
  
Return ONLY the email body text with proper line breaks, no subject line."""  
  
    response = client.chat.completions.create(  
        model=GPT_MODEL,  
        messages=[  
            {"role": "system", "content": "Generate a natural, personalized email body. Use proper line breaks (\\n\\n between paragraphs). Return only the email text."},  
            {"role": "user", "content": prompt}  
        ],  
        temperature=0.5  
    )  
      
    email_body = response.choices[0].message.content.strip()  
    signature = f"\n\n{candidate_name}"  
    if phone_number:  
        signature += f"\n{phone_number}"  
      
    return email_body + signature  
  
def create_eml_file(output_path: Path, email_body: str, cv_pdf: Path, cl_pdf: Path, subject: str):  
    from email.mime.multipart import MIMEMultipart  
    from email.mime.text import MIMEText  
    from email.mime.application import MIMEApplication  
      
    msg = MIMEMultipart()  
    msg['Subject'] = subject  
    msg['From'] = ""  
    msg['To'] = ""  
      
    msg.attach(MIMEText(email_body, 'plain', 'utf-8'))  
      
    with open(cv_pdf, 'rb') as f:  
        cv_attachment = MIMEApplication(f.read(), _subtype='pdf')  
        cv_attachment.add_header('Content-Disposition', 'attachment', filename=cv_pdf.name)  
        msg.attach(cv_attachment)  
      
    with open(cl_pdf, 'rb') as f:  
        cl_attachment = MIMEApplication(f.read(), _subtype='pdf')  
        cl_attachment.add_header('Content-Disposition', 'attachment', filename=cl_pdf.name)  
        msg.attach(cl_attachment)  
      
    with open(output_path, 'w') as f:  
        f.write(msg.as_string())  
  
def process_job_posting(job_posting: str, url: str = None) -> dict:  
    if not TEMPLATE_COVER_LETTER.exists():  
        print("Error: Template cover letter not found")  
        return {}  
      
    doc = Document(TEMPLATE_COVER_LETTER)  
    candidate_name = doc.paragraphs[0].text.strip()  
      
    language = detect_language(job_posting)  
    cover_letter_label = get_cover_letter_label(language)  
      
    company_name = get_company_name(job_posting)  
    position_title = get_position_title(job_posting)  
    job_location = get_job_location(job_posting)  
      
    print(f"Company: {company_name}, Position: {position_title}, Location: {job_location}, Language: {language}")  
      
    company_dir = OUTPUT_BASE_DIR / company_name  
    company_dir.mkdir(exist_ok=True, parents=True)  
      
    print("Getting CV updates...")  
    cv_updates = get_cv_updates(job_posting, language)  
      
    print("Generating cover letter...")  
    new_paragraphs = get_new_cover_letter_paragraphs(job_posting, company_name, TEMPLATE_COVER_LETTER, language)  
      
    output_cl_docx = get_unique_filename(company_dir / f"{candidate_name} {cover_letter_label}.docx")  
    modify_cover_letter(TEMPLATE_COVER_LETTER, output_cl_docx, new_paragraphs)  
      
    output_cl_pdf = output_cl_docx.with_suffix('.pdf')  
    convert_to_pdf_linux(output_cl_docx, output_cl_pdf)  
      
    if TEMPLATE_CV.exists():  
        output_cv_docx = get_unique_filename(company_dir / f"{candidate_name} CV.docx")  
        modify_cv_header(TEMPLATE_CV, output_cv_docx, cv_updates)  
          
        output_cv_pdf = output_cv_docx.with_suffix('.pdf')  
        convert_to_pdf_linux(output_cv_docx, output_cv_pdf)  
          
        print("Generating email...")  
        email_body = generate_email_body(company_name, position_title, language, TEMPLATE_COVER_LETTER)  
        email_subject = get_email_subject(language)  
        output_eml = get_unique_filename(company_dir / f"{candidate_name} Email.eml")  
        create_eml_file(output_eml, email_body, output_cv_pdf, output_cl_pdf, email_subject)  
          
        print(f"Files: {output_cl_docx.name}, {output_cl_pdf.name}, {output_cv_docx.name}, {output_cv_pdf.name}, {output_eml.name}\n")  
    else:  
        print(f"Files: {output_cl_docx.name}, {output_cl_pdf.name}\n")  
      
    return {  
        'company': company_name,  
        'position': position_title,  
        'location': job_location,  
        'language': language,  
        'date_processed': datetime.now().strftime("%Y-%m-%d %H:%M"),  
        'status': 'done'  
    }  
  
def main():      
    entries = read_csv_entries()  
    pending = [e for e in entries  
               if (e.get('status') or '').lower() != 'done'  
               and (e.get('url') or '').strip()]  
      
    if not pending:  
        print("No pending URLs in list.csv")  
        return  
      
    print(f"Found {len(pending)} pending URL(s)\n")  
      
    for i, entry in enumerate(pending, 1):  
        url = (entry.get('url') or '').strip()  
        if not url:  
            continue  
              
        print(f"[{i}/{len(pending)}] Fetching: {url}")  
        job_posting = fetch_job_posting_from_url(url)  
          
        if job_posting:  
            metadata = process_job_posting(job_posting, url)  
            if metadata:  
                update_csv_entry(url, metadata)  
        else:  
            print(f"Skipped: Failed to fetch {url}\n")  
            update_csv_entry(url, {'status': 'fetch_failed', 'date_processed': datetime.now().strftime("%Y-%m-%d %H:%M")})  
  
if __name__ == "__main__":  
    main()