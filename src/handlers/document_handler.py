import os
import shutil
from pathlib import Path
from docx import Document

from handlers.ai_handler import get_cv_replacements


def modify_cv_header(cv_path: Path, output_path: Path, updates: dict):
    shutil.copy(cv_path, output_path)
    doc = Document(output_path)

    header_template = os.getenv("CV_HEADER_TEMPLATE", "")
    if not header_template:
        doc.save(output_path)
        return

    replacements = get_cv_replacements(header_template, updates)
    if not replacements:
        doc.save(output_path)
        return

    new_header = header_template
    for old, new in replacements.items():
        if old and new:
            new_header = new_header.replace(old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if header_template in cell.text:
                    for para in cell.paragraphs:
                        full_text = "".join(r.text for r in para.runs)
                        if header_template in full_text:
                            new_full = full_text.replace(header_template, new_header)
                            for i, run in enumerate(para.runs):
                                if i == 0:
                                    run.text = new_full
                                else:
                                    run.text = ""
                            doc.save(output_path)
                            return

    for para in doc.paragraphs:
        full_text = "".join(r.text for r in para.runs)
        if header_template in full_text:
            new_full = full_text.replace(header_template, new_header)
            for i, run in enumerate(para.runs):
                if i == 0:
                    run.text = new_full
                else:
                    run.text = ""
            doc.save(output_path)
            return

    doc.save(output_path)


def modify_cover_letter(template_path: Path, output_path: Path, new_paragraphs: list):
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    for i, para in enumerate(doc.paragraphs):
        if i < len(new_paragraphs):
            for run in para.runs[1:]:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_paragraphs[i]
            else:
                para.text = new_paragraphs[i]

    doc.save(output_path)