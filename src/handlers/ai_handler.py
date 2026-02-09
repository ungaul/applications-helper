import os
import json
import re
from datetime import datetime
from openai import OpenAI
from docx import Document

def get_ai_client():
    return OpenAI(
        api_key=os.getenv("OPENAI_API_KEY"),
        base_url=os.getenv("AI_ENDPOINT", "https://api.openai.com/v1")
    )

def _chat(system: str, user: str, temperature: float = 0) -> str:
    client = get_ai_client()
    response = client.chat.completions.create(
        model=os.getenv("GPT_MODEL", "gpt-4.1-mini"),
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        temperature=temperature
    )
    return response.choices[0].message.content.strip()

def get_company_name(job_posting: str) -> str:
    return _chat(
        """Extract only the hiring company name. Respond with the name only, nothing else.
RULES:
- Return the company that is HIRING, not the job board or recruitment platform (e.g. ignore Indeed, LinkedIn, Welcome to the Jungle, Hellowork, etc.).
- The name must be a plausible company name: no numbers alone, no URLs, no gibberish.
- If you cannot confidently identify the company name, respond with exactly "Unknown".
- Double-check your answer: is this really the employer, not the platform hosting the ad?""",
        job_posting
    )

def get_position_title(job_posting: str) -> str:
    return _chat("Extract only the job position/title. Respond with the title only, nothing else.", job_posting)

def get_job_location(job_posting: str) -> str:
    return _chat(
        "Extract only the job location (city, country or remote). Respond with the location only, nothing else. If not specified, respond with 'Not specified'.",
        job_posting
    )

def detect_language(job_posting: str) -> str:
    return _chat(
        "Detect the language of the text and respond with only the ISO 639-1 language code (e.g., 'fr', 'en', 'de', 'es', 'ja', etc.). Return ONLY the 2-letter code, nothing else.",
        job_posting[:500]
    ).lower()

def get_cover_letter_label(language: str) -> str:
    return _chat(f"Return only the translation of 'Cover Letter' in the language with ISO code '{language}'. Return ONLY the translation, nothing else.", "Cover Letter")

def get_email_subject(position_title: str, job_field: str, candidate_name: str, language: str) -> str:
    return _chat(
        "Return ONLY a short professional email subject line for a job application. No quotes, no explanation.",
        f"Compose an email subject in language '{language}' for: position='{position_title}', field='{job_field}', candidate='{candidate_name}'. Format like: Candidature - [Position] - [Field] - [Name] (adapt 'Candidature' to the language).",
    )

def get_cv_updates(job_posting: str, language: str) -> dict:
    todays_date = datetime.now().strftime("%Y-%m-%d")
    prompt = f"""Extract CV header updates from this job posting.
JOB POSTING: {job_posting}
TODAY'S DATE: {todays_date}
Return JSON with:
- "job_field": The SECTOR/DOMAIN of the position, NOT the job title. Use 1-2 words maximum. Examples: Financial Analyst → "Finance", Management Controller → "Controlling". Keep it in the SAME LANGUAGE as the job posting (ISO code: {language}).
- "start_date": Start month in the format and language of the job posting. If not specified, use current month+1.
Return ONLY valid JSON."""

    content = _chat("Return only valid JSON. No other text.", prompt)
    content = content.replace('```json', '').replace('```', '').strip()

    try:
        return json.loads(content)
    except json.JSONDecodeError:
        print(f"Warning: Invalid JSON from GPT: {content}")
        return {"job_field": "Finance", "start_date": "February"}


def get_cv_replacements(header_text: str, updates: dict) -> dict:
    job_field = updates.get('job_field', '')
    start_date = updates.get('start_date', '')
    
    prompt = f"""Analyze this CV header text and identify EXACTLY which words need to be replaced.

CV HEADER TEXT:
{header_text}

REPLACEMENTS NEEDED:
- Find the business domain/field word (like "Finance", "Marketing", "Audit", "Controlling", etc.) and replace with: "{job_field}"
- Find the start month (like "Février", "March", "Januar", etc.) and replace with: "{start_date}"

Return a JSON object where keys are the EXACT original words found in the text, and values are what to replace them with.
Only include words that actually exist in the text and need replacing.
Example: {{"Finance": "Controlling", "Février": "Mars"}}

Return ONLY valid JSON, no explanations."""

    content = _chat("Return only valid JSON.", prompt)
    content = content.replace('```json', '').replace('```', '').strip()

    try:
        return json.loads(content)
    except json.JSONDecodeError:
        print(f"Warning: Invalid JSON for replacements: {content}")
        return {}


def get_new_cover_letter_paragraphs(job_posting: str, company_name: str, template_path, language: str) -> list:
    doc = Document(template_path)
    original_paragraphs = [p.text for p in doc.paragraphs]
    todays_date = datetime.now().strftime("%Y-%m-%d")

    prompt = f"""Return a JSON array of paragraph texts (strings) for a cover letter, matching the exact structure and number of paragraphs from the example. 

CRITICAL RULES:
1) Copy the ENTIRE example EXACTLY except for these specific changes:
   - Company address block (company name, department, city - do NOT invent street addresses)
   - Date
   - Job title in subject line and intro paragraph
   - Paragraph 3 content (see below)
   - Do NOT change the personal info (name, address, email, phone) present in the original template.
2) EVERYTHING ELSE must be IDENTICAL to the example - same personal info, same experiences, same wording, same closing.
3) The applicant's full name MUST appear once at the end of the letter, as in the original.
4) LANGUAGE QUALITY: Write in fluent, native-level {language}. 
   - Avoid anglicisms and literal translations.
   - Use natural expressions native speakers would use.
   - Avoid "patterns", "process", corporate and meaningless sentences.
5) After the subject line, there MUST be an empty paragraph (empty string "") before the salutation
6) The signature at the end should appear ONLY ONCE

JOB POSTING: {job_posting}
COMPANY: {company_name}
TODAY'S DATE: {todays_date}
ORIGINAL PARAGRAPHS:
{chr(10).join(f'{i}: "{p}"' for i, p in enumerate(original_paragraphs))}

Return a JSON array with EXACTLY {len(original_paragraphs)} elements, preserving empty strings where the original has them."""

    content = _chat("Return only a valid JSON array of strings. No other text.", prompt, temperature=0.5)
    content = content.replace('```json', '').replace('```', '').strip()

    try:
        return json.loads(content)
    except json.JSONDecodeError:
        content = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', content)
        try:
            return json.loads(content)
        except json.JSONDecodeError:
            print("Failed to parse. Returning original paragraphs.")
            return original_paragraphs


def generate_email_body(company_name: str, position_title: str, language: str, template_path) -> str:
    doc = Document(template_path)
    candidate_name = doc.paragraphs[0].text.strip()

    phone_number = ""
    for para in doc.paragraphs[:5]:
        phone_match = re.search(r'\+?\d[\d\s\-\(\)]{8,}', para.text)
        if phone_match:
            phone_number = phone_match.group(0)
            break

    prompt = f"""Generate a professional, natural email body for a job application.
COMPANY: {company_name}
POSITION: {position_title}
LANGUAGE: ISO code {language}

RULES:
- Two short paragraphs: first to introduce yourself and mention the position, second to mention attachments and availability.
- Use the CORRECT gender for the candidate name "{candidate_name}" (masculine/feminine forms, no gender-neutral parentheses like "motivé(e)").
- Sound natural and human, not corporate or robotic. Write like a real person would.
- No filler sentences, no meaningless corporate fluff.
- End with just the closing phrase (e.g. "Cordialement,"), NO name after it.
Return ONLY the email body text."""

    email_body = _chat("Generate a natural, human-sounding email body. Return only the email text.", prompt, temperature=0.5)
    signature = f"\n\n{candidate_name}"
    if phone_number:
        signature += f"\n{phone_number}"

    return email_body + signature